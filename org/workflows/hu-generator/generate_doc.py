#!/usr/bin/env python3
"""Gera o .docx de HU/HT no padrao do projeto a partir do .md consolidado.

Quando a secao 8 possui headings de prints, insere automaticamente as imagens de
`prototipo-prints/<IDENTIFICACAO>/` (ou `prototipo-prints/` no formato legado),
agrupadas pelo prefixo numerico (04a/04b = uma print logica).

Uso:  python3 generate_doc.py <md_path> <docx_out> [--label "HISTORIA TECNICA"]

A lib e python-docx (leve, pure-python). Layout identico ao gerador antigo em Node/docx.
O rotulo do header e inferido do frontmatter (`tipo: HU|HT`); --label sobrescreve.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

# ---- dimensoes (twips do gerador original) ----
TABLE_LABEL_W = 2220
TABLE_VAL_W = 8246
GREEN = RGBColor(0x38, 0x76, 0x1D)
DARK = RGBColor(0x1B, 0x1C, 0x1D)
LOGO = Path(__file__).resolve().parent / "assets" / "header_logo.png"
PROTOTYPE_PRINT_RE = re.compile(r"^(\d+)([a-z]?)_.+\.png$", re.IGNORECASE)
DOCUMENT_ID_RE = re.compile(r"\b(?:HU|HT)\d+(?:\.\d+)*\b", re.IGNORECASE)

ANCHOR_XML = (
    '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
    ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
    '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"'
    ' behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="column"><wp:posOffset>-18409</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>19050</wp:posOffset></wp:positionV>'
    '<wp:extent cx="{cx}" cy="{cy}"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
    '<wp:docPr id="7" name="header_logo"/><wp:cNvGraphicFramePr/>'
    '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
    '<pic:pic><pic:nvPicPr><pic:cNvPr id="7" name="header_logo"/><pic:cNvPicPr/></pic:nvPicPr>'
    '<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
    '<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
    '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
    '</pic:pic></a:graphicData></a:graphic></wp:anchor></w:drawing>'
)


# ---------------- helpers de formatacao ----------------
def runs_from_md(p, text, size=12, color=None):
    """Adiciona runs a um paragrafo, interpretando **negrito**."""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = i % 2 == 1  # segmentos impares vieram de **...**
        if color is not None:
            r.font.color.rgb = color


def body_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    runs_from_md(p, text)
    return p


def bullet_base(doc):
    # bullet nativo do Word (estilo "List Bullet" -> numbering.xml), sem glifo manual
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)  # igual ao espacamento dos paragrafos (body_para)
    return p


def fix_bullet_glyph(doc):
    """Ajusta o estilo nativo 'List Bullet' do template padrao do python-docx, que vem
    com 3 problemas pro padrao do projeto:
    1. Bullet em fonte Symbol (glifo privado) -> exibe caractere estranho fora do Word/Windows.
       Troca pro bullet padrao (U+2022) em Arial.
    2. Recuo padrao do template (360/360, bullet flush na borda) nao bate com o padrao
       Google Docs (bullet a 0.25" da borda, texto a 0.5"). Ajusta pra left=720 hanging=360.
    3. `contextualSpacing` zera o `space_after` entre bullets consecutivos, deixando o
       espacamento entre bullets menor que o espacamento entre paragrafos de texto. Remove.
    """
    numbering_part = doc.part.numbering_part
    for lvl in numbering_part.element.findall(
        ".//" + qn("w:abstractNum") + "/" + qn("w:lvl")
    ):
        pstyle = lvl.find(qn("w:pStyle"))
        if pstyle is None or pstyle.get(qn("w:val")) != "ListBullet":
            continue
        lvl_text = lvl.find(qn("w:lvlText"))
        if lvl_text is not None:
            lvl_text.set(qn("w:val"), "•")
        rpr = lvl.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                rfonts.set(qn("w:ascii"), "Arial")
                rfonts.set(qn("w:hAnsi"), "Arial")
        ppr = lvl.find(qn("w:pPr"))
        if ppr is not None:
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                ind.set(qn("w:left"), "720")
                ind.set(qn("w:hanging"), "360")
            tab = ppr.find(qn("w:tabs") + "/" + qn("w:tab"))
            if tab is not None:
                tab.set(qn("w:pos"), "720")

    list_bullet_style = doc.styles["List Bullet"].element
    ppr = list_bullet_style.find(qn("w:pPr"))
    if ppr is not None:
        cs = ppr.find(qn("w:contextualSpacing"))
        if cs is not None:
            ppr.remove(cs)


def bullet_para(doc, text):
    p = bullet_base(doc)
    runs_from_md(p, text)


def ca_para(doc, text):
    """Bullet de CA: codigo verde, palavras-chave em negrito escuro."""
    p = bullet_base(doc)
    m = re.match(r"^(\*\*CA\d+:\*\*)\s+(.+)$", text)
    if not m:
        runs_from_md(p, text)
        return
    r = p.add_run(m.group(1).replace("**", "") + " ")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = GREEN
    rest = m.group(2)
    for i, seg in enumerate(re.split(r"(\*\*Dado que\*\*|\*\*Quando\*\*|\*\*Então\*\*)", rest)):
        if seg == "":
            continue
        kw = seg.startswith("**")
        rr = p.add_run(seg.replace("**", "") if kw else seg)
        rr.font.name = "Arial"
        rr.font.size = Pt(12)
        rr.bold = kw
        rr.font.color.rgb = DARK


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "8")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)


def add_row_table(doc, label, value):
    """Cada linha vira uma tabela de 1 linha x 2 colunas (igual ao gerador original)."""
    t = doc.add_table(rows=1, cols=2)
    t.allow_autofit = False
    set_table_borders(t)
    cells = t.rows[0].cells
    cells[0].width = Pt(TABLE_LABEL_W / 20)
    cells[1].width = Pt(TABLE_VAL_W / 20)
    lr = cells[0].paragraphs[0].add_run(label)
    lr.font.name = "Arial"
    lr.font.size = Pt(12)
    lr.bold = True
    runs_from_md(cells[1].paragraphs[0], value)


def add_field(run, instr, placeholder="1"):
    r = run._r
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    ft = OxmlElement("w:t")
    ft.text = placeholder
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    for el in (fb, it, fs, ft, fe):
        r.append(el)


def add_floating_logo(paragraph):
    if not LOGO.exists():
        return
    rid, _ = paragraph.part.get_or_add_image(str(LOGO))
    run = paragraph.add_run()
    run._r.append(parse_xml(ANCHOR_XML.format(cx=319 * 9525, cy=26 * 9525, rid=rid)))


def prototype_prints_dir(md_path, metadata):
    """Resolve a pasta de prints da HU/HT, preservando o formato legado."""
    base_dir = Path(md_path).resolve().parent / "prototipo-prints"
    document_id = None
    for label, value in metadata:
        if label in ("Identificação da HU", "Identificação da HT"):
            match = DOCUMENT_ID_RE.search(value)
            if match:
                document_id = match.group(0).upper()
                break

    if document_id:
        document_dir = base_dir / document_id
        if document_dir.is_dir():
            return document_dir

    return base_dir


def collect_prototype_prints(md_path, metadata, sections):
    """Valida e agrupa prints da secao 8 na mesma ordem dos headings H4."""
    prototype_headings = [
        payload[1]
        for section in sections
        if section["heading"].startswith("8.")
        for block_type, payload in section["blocks"]
        if block_type == "subheading" and payload[0] == 4
    ]
    prints_dir = prototype_prints_dir(md_path, metadata)

    if not prototype_headings:
        return []
    if not prints_dir.is_dir():
        raise ValueError(
            f"A secao 8 possui {len(prototype_headings)} titulo(s), mas a pasta "
            f"de prints nao existe: {prints_dir}"
        )

    groups = {}
    for image_path in sorted(prints_dir.glob("*.png")):
        match = PROTOTYPE_PRINT_RE.match(image_path.name)
        if not match:
            raise ValueError(
                f"Nome de print invalido: {image_path.name}. "
                "Use NN_descricao.png ou NNa_descricao.png."
            )
        number = int(match.group(1))
        suffix = match.group(2).lower()
        groups.setdefault(number, []).append((suffix, image_path))

    expected_numbers = list(range(1, len(prototype_headings) + 1))
    if sorted(groups) != expected_numbers:
        found = ", ".join(f"{number:02d}" for number in sorted(groups)) or "nenhum"
        expected = ", ".join(f"{number:02d}" for number in expected_numbers)
        raise ValueError(
            f"Prints da secao 8 nao correspondem aos headings. "
            f"Esperado: {expected}. Encontrado: {found}."
        )

    result = []
    for number, heading in zip(expected_numbers, prototype_headings):
        parts = sorted(groups[number], key=lambda item: item[0])
        suffixes = [suffix for suffix, _ in parts]
        if "" in suffixes and len(suffixes) > 1:
            raise ValueError(
                f"A print {number:02d} mistura arquivo sem sufixo com partes a/b."
            )
        if suffixes and suffixes[0] != "":
            expected_suffixes = [chr(ord("a") + i) for i in range(len(suffixes))]
            if suffixes != expected_suffixes:
                raise ValueError(
                    f"Partes da print {number:02d} devem ser continuas a/b/c. "
                    f"Encontrado: {', '.join(suffixes)}."
                )
        result.append((heading, [path for _, path in parts]))
    return result


def add_prototype_prints(doc, heading, image_paths, max_width):
    """Insere uma print logica (uma ou mais partes) como imagens inline."""
    for index, image_path in enumerate(image_paths):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0 if index < len(image_paths) - 1 else 6)
        run = paragraph.add_run()
        shape = run.add_picture(str(image_path), width=Emu(max_width))
        part_label = "" if len(image_paths) == 1 else f" - parte {index + 1}"
        shape._inline.docPr.set("name", image_path.name)
        shape._inline.docPr.set("descr", f"{heading}{part_label}")


# ---------------- parser do .md ----------------
def parse_md(text):
    lines = text.split("\n")
    tipo = None
    metadata = []  # [(label, value)]
    sections = []  # [{heading, blocks:[(type, payload)]}]
    cur = None
    in_fm = False
    in_meta = False
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows and cur is not None:
            cur["blocks"].append(("table", table_rows))
        table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        # frontmatter
        if i == 0 and line.strip() == "---":
            in_fm = True
            i += 1
            continue
        if in_fm:
            if line.strip() == "---":
                in_fm = False
            else:
                m = re.match(r"tipo:\s*(\S+)", line)
                if m:
                    tipo = m.group(1).strip()
            i += 1
            continue

        if line.startswith("## Metadados"):
            in_meta = True
            i += 1
            continue
        if in_meta:
            m = re.match(r"- \*\*(.+?):\*\*\s*(.+)", line)
            if m:
                metadata.append((m.group(1).strip(), m.group(2).strip()))
                i += 1
                continue
            if line.startswith("#"):
                in_meta = False
            else:
                i += 1
                continue

        # cabecalho de secao (## N. ... ou ## Apendice ...)
        if re.match(r"^## (\d+\.|Apêndice)", line):
            flush_table()
            cur = {"heading": line[3:].strip(), "blocks": []}
            sections.append(cur)
            i += 1
            continue

        # subsecao ### (h3) ou #### (h4)
        m = re.match(r"^(#{3,4}) (.+)", line)
        if m and cur is not None:
            flush_table()
            level = len(m.group(1))  # 3 ou 4
            cur["blocks"].append(("subheading", (level, m.group(2).strip())))
            i += 1
            continue

        if cur is None:
            i += 1
            continue

        # linha de tabela | **label** | value |
        if re.match(r"^\|\s*\|\s*\|$", line) or re.match(r"^\|[\s\-|]+\|$", line):
            i += 1
            continue
        m = re.match(r"^\|\s*\*\*(.+?)\*\*\s*\|\s*(.+?)\s*\|$", line)
        if m:
            table_rows.append((m.group(1).strip(), m.group(2).strip()))
            i += 1
            continue
        else:
            flush_table()

        # CA
        if re.match(r"^- \*\*CA\d+:", line):
            cur["blocks"].append(("ca", line[2:].strip()))
            i += 1
            continue
        # bullet
        m = re.match(r"^- (.+)", line)
        if m:
            cur["blocks"].append(("bullet", m.group(1).strip()))
            i += 1
            continue
        # vazio
        if line.strip() == "":
            cur["blocks"].append(("empty", None))
            i += 1
            continue
        # texto comum
        if not line.startswith(("#", "|", "-")):
            cur["blocks"].append(("paragraph", line.strip()))
        i += 1

    flush_table()
    return tipo, metadata, sections


# ---------------- estilos + montagem ----------------
def setup_styles(doc):
    fix_bullet_glyph(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    for hid, sz in (("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)):
        st = doc.styles[hid]
        st.font.name = "Arial"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.line_spacing = 1.5


def build(md_path, out_path, label=None):
    tipo, metadata, sections = parse_md(Path(md_path).read_text(encoding="utf-8"))
    prototype_prints = collect_prototype_prints(md_path, metadata, sections)
    prototype_print_index = 0
    if label is None:
        label = "HISTÓRIA TÉCNICA" if (tipo or "").upper() == "HT" else "HISTÓRIA DE USUÁRIO"

    doc = Document()
    setup_styles(doc)

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Pt(36)  # 720 twips
    sec.header_distance = Pt(35.4)
    sec.footer_distance = Pt(35.4)
    prototype_max_width = int(sec.page_width - sec.left_margin - sec.right_margin)

    # header: logo flutuante (esq) + label (dir) no mesmo paragrafo
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.line_spacing = 2.0
    hp.paragraph_format.right_indent = Pt(1.5)
    tr = hp.add_run(label)
    tr.font.name = "Calibri"
    tr.font.size = Pt(18)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0, 0, 0)
    add_floating_logo(hp)
    sec.header.add_paragraph()

    # footer: numero de pagina centralizado
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = fp.add_run()
    pr.font.name = "Aptos"
    pr.font.small_caps = True
    pr.font.color.rgb = RGBColor(0x15, 0x60, 0x82)
    add_field(pr, "PAGE")

    # 5 linhas em branco
    for _ in range(5):
        doc.add_paragraph()

    # metadados
    for lbl, val in metadata:
        p = doc.add_paragraph()
        a = p.add_run(f"{lbl}: ")
        a.font.name = "Arial"
        a.font.size = Pt(17)
        a.bold = True
        b = p.add_run(val)
        b.font.name = "Arial"
        b.font.size = Pt(17)
    doc.add_paragraph()

    # sumario (TOC real)
    doc.add_paragraph("Sumário", style="Heading 1")
    tp = doc.add_paragraph()
    add_field(tp.add_run(), 'TOC \\o "1-3" \\h \\z \\u', placeholder="Atualize o sumário (F9).")
    doc.add_paragraph()

    # secoes
    for s in sections:
        # apendices (trilha de discovery, GLs a copiar pro Drive) NAO vao pro .docx
        if s["heading"].startswith("Apêndice"):
            continue
        doc.add_paragraph(s["heading"], style="Heading 1")
        # Todas as secoes renderizam o conteudo completo do .md (autocontido): RN em SBVR,
        # MSG com texto literal e GL como bullets sao bullets comuns; CA tem estilo proprio.
        for btype, payload in s["blocks"]:
            if btype == "paragraph":
                body_para(doc, payload)
            elif btype == "subheading":
                level, txt = payload
                heading_paragraph = doc.add_paragraph(
                    txt, style="Heading 3" if level == 4 else "Heading 2"
                )
                if (
                    s["heading"].startswith("8.")
                    and level == 4
                    and prototype_prints
                ):
                    expected_heading, image_paths = prototype_prints[prototype_print_index]
                    if txt != expected_heading:
                        raise ValueError(
                            f"Titulo da print {prototype_print_index + 1:02d} mudou durante a geracao: "
                            f"esperado '{expected_heading}', encontrado '{txt}'."
                        )
                    heading_paragraph.paragraph_format.keep_with_next = True
                    add_prototype_prints(doc, txt, image_paths, prototype_max_width)
                    prototype_print_index += 1
            elif btype == "table":
                for lbl, val in payload:
                    add_row_table(doc, lbl, val)
            elif btype == "ca":
                ca_para(doc, payload)
            elif btype == "bullet":
                bullet_para(doc, payload)
            elif btype == "empty":
                doc.add_paragraph()

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    image_count = sum(len(paths) for _, paths in prototype_prints)
    print(f"Generated: {out_path} ({image_count} prototype image(s))")


def main(argv):
    label = None
    args = []
    it = iter(argv)
    for a in it:
        if a == "--label":
            label = next(it)
        else:
            args.append(a)
    if len(args) < 2:
        print('Usage: python3 generate_doc.py <md_path> <docx_out> [--label "..."]', file=sys.stderr)
        return 1
    build(args[0], args[1], label)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
